"""DOCX export for transcripts using python-docx."""

from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from src.core.database import Recording, Summary, Transcript, TranscriptSegment
from src.core.recording_titles import display_recording_title


def export_docx(
    recording: Recording,
    transcript: Optional[Transcript],
    segments: list[TranscriptSegment],
    summary: Optional[Summary],
    output_path: Path,
) -> None:
    """
    Export transcript to DOCX file.

    Args:
        recording: Recording data
        transcript: Transcript data
        segments: List of transcript segments
        summary: Summary data
        output_path: Path to output file
    """
    doc = Document()

    # Title
    title = doc.add_heading(display_recording_title(recording), level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Info section
    info_para = doc.add_paragraph()

    if recording.created_at:
        if isinstance(recording.created_at, str):
            date_str = recording.created_at
        else:
            date_str = recording.created_at.strftime("%d.%m.%Y %H:%M")
        info_para.add_run(f"Дата: {date_str}\n")

    if recording.duration_seconds:
        minutes = recording.duration_seconds // 60
        seconds = recording.duration_seconds % 60
        info_para.add_run(f"Длительность: {minutes}:{seconds:02d}\n")

    if transcript and transcript.language:
        lang_names = {"ru": "Русский", "en": "English"}
        lang = lang_names.get(transcript.language, transcript.language)
        info_para.add_run(f"Язык: {lang}\n")

    # Summary section
    if summary:
        doc.add_heading("Краткое содержание", level=1)

        if summary.summary:
            doc.add_paragraph(summary.summary)

        if summary.key_points:
            doc.add_heading("Ключевые темы", level=2)
            for point in summary.key_points:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(point)

        if summary.decisions:
            doc.add_heading("Принятые решения", level=2)
            for decision in summary.decisions:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"✓ {decision}")

        if summary.action_items:
            doc.add_heading("Задачи", level=2)
            for item in summary.action_items:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"☐ {item}")

    # Transcript section
    doc.add_heading("Транскрипт", level=1)

    if segments:
        current_speaker = None
        for seg in segments:
            speaker_label = seg.display_speaker
            if speaker_label != current_speaker:
                current_speaker = speaker_label

                # Speaker header
                speaker_para = doc.add_paragraph()
                speaker_run = speaker_para.add_run(f"\n{speaker_label}:")
                speaker_run.bold = True
                speaker_run.font.size = Pt(12)
                speaker_run.font.color.rgb = RGBColor(0, 100, 0)

            # Segment text with timestamp
            start_min = int(seg.start_time // 60)
            start_sec = int(seg.start_time % 60)
            timestamp = f"[{start_min}:{start_sec:02d}]"

            para = doc.add_paragraph()

            # Timestamp in gray
            ts_run = para.add_run(timestamp + " ")
            ts_run.font.color.rgb = RGBColor(128, 128, 128)
            ts_run.font.size = Pt(10)

            # Text
            para.add_run(seg.text)

    elif transcript and transcript.full_text:
        doc.add_paragraph(transcript.full_text)
    else:
        doc.add_paragraph("(Транскрипт отсутствует)")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer.add_run("Сгенерировано Meeting Note")
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.font.size = Pt(10)

    # Save document
    doc.save(str(output_path))
